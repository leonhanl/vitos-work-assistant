from io import BytesIO

import pytest
from docx import Document

from m365_mcp.document_parser import DocumentParseError, extract_docx_text, extract_text


def test_extract_docx_text() -> None:
    document = Document()
    document.add_heading("VPN Guide", level=1)
    document.add_paragraph("Install the corporate VPN client.")
    table = document.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Platform"
    table.cell(0, 1).text = "Windows"
    buffer = BytesIO()
    document.save(buffer)

    text = extract_docx_text(buffer.getvalue())

    assert "VPN Guide" in text
    assert "Install the corporate VPN client." in text
    assert "Platform\tWindows" in text


def test_pdf_is_explicitly_unsupported() -> None:
    with pytest.raises(DocumentParseError, match="PDF parsing is not supported"):
        extract_text("policy.pdf", b"%PDF")
