"""Text extraction for the intentionally small set of supported file types."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document


class DocumentParseError(ValueError):
    """A clear, safe-to-display document parsing error."""


def validate_supported_type(name: str) -> None:
    """Raise a specific error unless the file has a supported extension."""
    extension = Path(name).suffix.lower()
    if extension in {".docx", ".txt", ".md"}:
        return
    if extension == ".pdf":
        raise DocumentParseError(
            "PDF parsing is not supported in the current version."
        )
    raise DocumentParseError(
        f"Unsupported file type: {extension or '(no extension)'}. "
        "Supported types are .docx, .txt, and .md."
    )


def extract_text(name: str, content: bytes) -> str:
    """Extract text from a DOCX, UTF-8 text file, or Markdown file."""
    validate_supported_type(name)
    extension = Path(name).suffix.lower()
    if extension == ".docx":
        return extract_docx_text(content)
    if extension in {".txt", ".md"}:
        try:
            return content.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise DocumentParseError(
                f"{extension} files must be encoded as UTF-8 in the current version."
            ) from exc
    raise AssertionError("validated file type was not handled")


def extract_docx_text(content: bytes) -> str:
    """Extract non-empty paragraphs and table rows from a DOCX byte string."""
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise DocumentParseError("The DOCX file could not be parsed.") from exc

    lines = [paragraph.text.strip() for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                lines.append("\t".join(cells))
    return "\n".join(line for line in lines if line)
